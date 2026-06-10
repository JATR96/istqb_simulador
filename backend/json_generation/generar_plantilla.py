import csv

from docx import Document

from docx.shared import RGBColor


# ==========================================
# CONFIG
# ==========================================

CERTIFICATION = "Foundation Tester"

CSV_ENTRADA = "metadata.csv"

WORD_SALIDA = f"plantilla_{CERTIFICATION}.docx"


# ==========================================
# COLORES
# ==========================================

VERDE = RGBColor(0, 128, 0)

ROJO = RGBColor(255, 0, 0)


# ==========================================
# HELPERS
# ==========================================

def agregar_linea_verde(
        doc,
        texto
):

    p = doc.add_paragraph()

    run = p.add_run(texto)

    run.font.color.rgb = VERDE

    run.bold = True


def agregar_linea_roja(
        doc,
        texto
):

    p = doc.add_paragraph()

    run = p.add_run(texto)

    run.font.color.rgb = ROJO

    run.bold = True


def agregar_linea_normal(
        doc,
        texto=""
):

    doc.add_paragraph(texto)

# ==========================================
# OBTENER CANTIDAD RESPUESTAS
# ==========================================

def obtener_cantidad_respuestas(respuestas_correctas):
    """
    Ejemplos:
    a -> 1
    a,e -> 2
    a,c,e -> 3
    """

    respuestas_correctas = respuestas_correctas.strip()

    if not respuestas_correctas:
        return 0

    return len([
        x.strip()
        for x in respuestas_correctas.split(",")
        if x.strip()
    ])

# ==========================================
# GENERADOR
# ==========================================

def generar_word():

    doc = Document()

    with open(
            CSV_ENTRADA,
            encoding="utf-8"
    ) as archivo:

        reader = csv.DictReader(
            archivo
        )

        for fila in reader:

            cantidad_respuestas = obtener_cantidad_respuestas(
                fila["respuestas_correctas"]
            )

            agregar_linea_normal(
                doc,
                "=== QUESTION START ==="
            )

            agregar_linea_verde(
                doc,
                f"CERTIFICATION: {CERTIFICATION}"
            )

            agregar_linea_verde(
                doc,
                f"LEARNING_OBJECTIVE: "
                f"{fila['learning_objective']}"
            )

            agregar_linea_verde(
                doc,
                f"K_LEVEL: "
                f"{fila['k_level']}"
            )

            agregar_linea_verde(
                doc,
                f"POINTS: "
                f"{fila['points']}"
            )

            agregar_linea_verde(
                doc,
                f"CANTIDAD_RESPUESTAS: "
                f"{cantidad_respuestas}"
            )

            agregar_linea_verde(
                doc,
                f"RESPUESTAS_CORRECTAS: "
                f"{fila['respuestas_correctas']}"
            )

            agregar_linea_normal(
                doc,
                "[es]"
            )

            agregar_linea_roja(
                doc,
                "PREGUNTA:"
            )

            agregar_linea_normal(doc)
            agregar_linea_normal(doc)
            agregar_linea_normal(doc)

            agregar_linea_roja(
                doc,
                "IMAGES:"
            )

            agregar_linea_normal(doc)
            agregar_linea_normal(doc)

            agregar_linea_roja(
                doc,
                "OPCIONES:"
            )

            agregar_linea_normal(doc)
            agregar_linea_normal(doc)
            agregar_linea_normal(doc)

            agregar_linea_roja(
                doc,
                "EXPLICACION:"
            )

            agregar_linea_normal(doc)
            agregar_linea_normal(doc)
            agregar_linea_normal(doc)

            agregar_linea_normal(
                doc,
                "[/es]"
            )

            agregar_linea_normal(doc)

            agregar_linea_normal(
                doc,
                "[en]"
            )

            agregar_linea_roja(
                doc,
                "QUESTION:"
            )

            agregar_linea_normal(doc)
            agregar_linea_normal(doc)
            agregar_linea_normal(doc)

            agregar_linea_roja(
                doc,
                "IMAGES:"
            )

            agregar_linea_normal(doc)
            agregar_linea_normal(doc)

            agregar_linea_roja(
                doc,
                "OPTIONS:"
            )

            agregar_linea_normal(doc)
            agregar_linea_normal(doc)
            agregar_linea_normal(doc)

            agregar_linea_roja(
                doc,
                "EXPLANATION:"
            )

            agregar_linea_normal(doc)
            agregar_linea_normal(doc)
            agregar_linea_normal(doc)

            agregar_linea_normal(
                doc,
                "[/en]"
            )

            agregar_linea_normal(
                doc,
                "=== QUESTION END ==="
            )

            doc.add_page_break()

    doc.save(
        WORD_SALIDA
    )

    print(
        f"\nWord generado:"
    )

    print(
        WORD_SALIDA
    )


# ==========================================
# MAIN
# ==========================================

if __name__ == "__main__":

    generar_word()